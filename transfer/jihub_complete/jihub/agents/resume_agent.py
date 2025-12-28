#!/usr/bin/env python3
"""
ResumeAgent - JIHUB Resume Transformation Engine
=================================================
AI-powered resume transformation preserving user-selected attributes.
Uses hybrid template system: extract format → transform content → merge back.

Usage:
    python resume_agent.py transform --input resume.pdf --type content_swap --output new_resume.pdf
    python resume_agent.py parse --input resume.docx --output template.json
    python resume_agent.py interactive
"""

import os
import sys
import json
import hashlib
import argparse
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, asdict, field
from typing import List, Dict, Optional, Any
from enum import Enum
import re

# Document parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[!] pdfplumber not installed - PDF parsing disabled")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[!] python-docx not installed - DOCX parsing disabled")

# Gemini API
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("[!] google-generativeai not installed - Gemini integration disabled")


class TransformType(Enum):
    """Types of resume transformations."""
    CONTENT_SWAP = "content_swap"      # Keep style/format/structure, change content
    FORMAT_SWAP = "format_swap"        # Keep content, change format/style/structure
    STYLE_SWAP = "style_swap"          # Keep content/format, change style
    STRUCTURE_SWAP = "structure_swap"  # Keep content/style, change structure
    HYBRID = "hybrid"                  # User-defined combination


@dataclass
class ResumeSection:
    """Represents a section of a resume."""
    section_type: str  # e.g., 'header', 'experience', 'education', 'skills'
    title: str
    content: str
    order: int
    formatting: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self):
        return asdict(self)


@dataclass
class ResumeTemplate:
    """Template representation of a resume preserving formatting."""
    original_path: str
    original_hash: str
    file_type: str  # 'pdf' or 'docx'
    sections: List[ResumeSection]
    global_formatting: Dict[str, Any]
    extracted_at: str
    
    # Style attributes
    tone: str = "professional"  # professional, casual, technical
    bullet_style: str = "bullet"  # bullet, dash, number, none
    verb_tense: str = "past"  # past, present, mixed
    
    def to_dict(self):
        d = asdict(self)
        d['sections'] = [s.to_dict() for s in self.sections]
        return d
    
    @classmethod
    def from_dict(cls, data: dict) -> 'ResumeTemplate':
        sections = [ResumeSection(**s) for s in data.pop('sections', [])]
        return cls(sections=sections, **data)


@dataclass
class TransformResult:
    """Result of a transformation operation."""
    success: bool
    input_path: str
    output_path: Optional[str]
    transform_type: str
    template_used: Optional[ResumeTemplate]
    gemini_response: Optional[Dict]
    format_preservation_score: float
    error_message: Optional[str] = None
    processing_time_ms: int = 0
    
    def to_dict(self):
        d = asdict(self)
        if self.template_used:
            d['template_used'] = self.template_used.to_dict()
        return d


class ResumeParser:
    """Parses resumes and extracts template information."""
    
    SECTION_KEYWORDS = {
        'experience': ['experience', 'work history', 'employment', 'professional experience', 'work experience'],
        'education': ['education', 'academic', 'degrees', 'certifications', 'training'],
        'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'proficiencies'],
        'summary': ['summary', 'objective', 'profile', 'about', 'overview'],
        'projects': ['projects', 'portfolio', 'achievements', 'accomplishments'],
        'contact': ['contact', 'email', 'phone', 'address', 'linkedin'],
        'header': []  # Detected by position, not keywords
    }
    
    def __init__(self):
        self.current_file = None
    
    def _compute_hash(self, filepath: str) -> str:
        """Compute MD5 hash of file."""
        with open(filepath, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def _detect_section_type(self, text: str) -> str:
        """Detect section type from text content."""
        text_lower = text.lower().strip()
        
        for section_type, keywords in self.SECTION_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower[:50]:  # Check first 50 chars
                    return section_type
        
        return 'other'
    
    def _extract_formatting_pdf(self, page) -> Dict[str, Any]:
        """Extract formatting info from PDF page."""
        formatting = {
            'page_width': page.width,
            'page_height': page.height,
            'has_tables': len(page.extract_tables()) > 0,
            'char_count': len(page.extract_text() or ''),
        }
        
        # Try to extract font info
        try:
            chars = page.chars
            if chars:
                fonts = set(c.get('fontname', 'unknown') for c in chars[:100])
                sizes = set(c.get('size', 12) for c in chars[:100])
                formatting['fonts'] = list(fonts)
                formatting['font_sizes'] = list(sizes)
        except:
            pass
        
        return formatting
    
    def _extract_formatting_docx(self, doc: Document) -> Dict[str, Any]:
        """Extract formatting info from DOCX document."""
        formatting = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections),
        }
        
        # Extract style info
        styles = set()
        fonts = set()
        sizes = set()
        
        for para in doc.paragraphs[:20]:  # Sample first 20 paragraphs
            if para.style:
                styles.add(para.style.name)
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    sizes.add(run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size)
        
        formatting['styles'] = list(styles)
        formatting['fonts'] = list(fonts)
        formatting['font_sizes'] = list(sizes)
        
        return formatting
    
    def parse_pdf(self, filepath: str) -> ResumeTemplate:
        """Parse a PDF resume."""
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber required for PDF parsing")
        
        sections = []
        global_formatting = {}
        
        with pdfplumber.open(filepath) as pdf:
            full_text = ""
            
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                full_text += text + "\n"
                
                if i == 0:
                    global_formatting = self._extract_formatting_pdf(page)
            
            # Split into sections
            lines = full_text.split('\n')
            current_section = []
            current_type = 'header'
            order = 0
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this line starts a new section
                detected_type = self._detect_section_type(line)
                
                if detected_type != 'other' and detected_type != current_type:
                    # Save previous section
                    if current_section:
                        sections.append(ResumeSection(
                            section_type=current_type,
                            title=current_section[0] if current_section else "",
                            content='\n'.join(current_section),
                            order=order,
                            formatting={}
                        ))
                        order += 1
                    
                    current_section = [line]
                    current_type = detected_type
                else:
                    current_section.append(line)
            
            # Don't forget last section
            if current_section:
                sections.append(ResumeSection(
                    section_type=current_type,
                    title=current_section[0] if current_section else "",
                    content='\n'.join(current_section),
                    order=order,
                    formatting={}
                ))
        
        return ResumeTemplate(
            original_path=filepath,
            original_hash=self._compute_hash(filepath),
            file_type='pdf',
            sections=sections,
            global_formatting=global_formatting,
            extracted_at=datetime.now().isoformat()
        )
    
    def parse_docx(self, filepath: str) -> ResumeTemplate:
        """Parse a DOCX resume."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX parsing")
        
        doc = Document(filepath)
        sections = []
        global_formatting = self._extract_formatting_docx(doc)
        
        current_section = []
        current_type = 'header'
        order = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Extract paragraph formatting
            para_formatting = {
                'style': para.style.name if para.style else None,
                'alignment': str(para.alignment) if para.alignment else None,
            }
            
            # Check if this starts a new section
            detected_type = self._detect_section_type(text)
            
            if detected_type != 'other' and detected_type != current_type:
                # Save previous section
                if current_section:
                    sections.append(ResumeSection(
                        section_type=current_type,
                        title=current_section[0]['text'] if current_section else "",
                        content='\n'.join(item['text'] for item in current_section),
                        order=order,
                        formatting=current_section[0].get('formatting', {}) if current_section else {}
                    ))
                    order += 1
                
                current_section = [{'text': text, 'formatting': para_formatting}]
                current_type = detected_type
            else:
                current_section.append({'text': text, 'formatting': para_formatting})
        
        # Don't forget last section
        if current_section:
            sections.append(ResumeSection(
                section_type=current_type,
                title=current_section[0]['text'] if current_section else "",
                content='\n'.join(item['text'] for item in current_section),
                order=order,
                formatting=current_section[0].get('formatting', {}) if current_section else {}
            ))
        
        return ResumeTemplate(
            original_path=filepath,
            original_hash=self._compute_hash(filepath),
            file_type='docx',
            sections=sections,
            global_formatting=global_formatting,
            extracted_at=datetime.now().isoformat()
        )
    
    def parse(self, filepath: str) -> ResumeTemplate:
        """Parse a resume file (auto-detect type)."""
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        ext = filepath.suffix.lower()
        
        if ext == '.pdf':
            return self.parse_pdf(str(filepath))
        elif ext in ['.docx', '.doc']:
            return self.parse_docx(str(filepath))
        else:
            raise ValueError(f"Unsupported file type: {ext}")


class GeminiTransformer:
    """Handles content transformation via Gemini API."""
    
    def __init__(self, api_key: Optional[str] = None):
        if not GEMINI_AVAILABLE:
            raise ImportError("google-generativeai required for Gemini integration")
        
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY not set. Set environment variable or pass api_key.")
        
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel('gemini-pro')
    
    def _build_transform_prompt(
        self,
        template: ResumeTemplate,
        transform_type: TransformType,
        instructions: Optional[str] = None
    ) -> str:
        """Build the transformation prompt for Gemini."""
        
        sections_json = json.dumps([s.to_dict() for s in template.sections], indent=2)
        
        base_prompt = f"""You are a professional resume transformation assistant.

CURRENT RESUME STRUCTURE:
{sections_json}

TRANSFORMATION TYPE: {transform_type.value}

TRANSFORMATION RULES:
"""
        
        if transform_type == TransformType.CONTENT_SWAP:
            base_prompt += """
- PRESERVE: All formatting, structure, section order, bullet styles, tone
- CHANGE: The actual content (job titles, company names, descriptions, skills)
- Generate realistic but fictional content that matches the resume's style
- Maintain the same number of items in each section
"""
        
        elif transform_type == TransformType.STYLE_SWAP:
            base_prompt += """
- PRESERVE: Content meaning, format layout, structure
- CHANGE: Tone (professional/casual/technical), verb choices, sentence structure
- Transform bullet points to paragraphs or vice versa if appropriate
- Adjust formality level while keeping the same information
"""
        
        elif transform_type == TransformType.STRUCTURE_SWAP:
            base_prompt += """
- PRESERVE: Content, style/tone
- CHANGE: Section order, hierarchy, grouping
- Reorganize sections (e.g., put Education before Experience)
- May combine or split sections logically
"""
        
        elif transform_type == TransformType.FORMAT_SWAP:
            base_prompt += """
- PRESERVE: Content only
- CHANGE: Everything else - style, structure, formatting suggestions
- Provide a complete restructuring with new organization
"""
        
        if instructions:
            base_prompt += f"\n\nADDITIONAL INSTRUCTIONS:\n{instructions}\n"
        
        base_prompt += """

OUTPUT FORMAT:
Return a JSON object with this exact structure:
{
  "sections": [
    {
      "section_type": "header|experience|education|skills|summary|projects|other",
      "title": "Section Title",
      "content": "Full section content...",
      "order": 0
    }
  ],
  "transformation_notes": "Brief explanation of changes made",
  "quality_score": 0.0-1.0,
  "format_preservation_score": 0.0-1.0
}

Return ONLY valid JSON, no markdown formatting or explanations outside the JSON.
"""
        
        return base_prompt
    
    def transform(
        self,
        template: ResumeTemplate,
        transform_type: TransformType,
        instructions: Optional[str] = None
    ) -> Dict[str, Any]:
        """Transform resume content via Gemini."""
        
        prompt = self._build_transform_prompt(template, transform_type, instructions)
        
        try:
            response = self.model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Clean up response (remove markdown code blocks if present)
            if response_text.startswith('```'):
                response_text = re.sub(r'^```json?\n?', '', response_text)
                response_text = re.sub(r'\n?```$', '', response_text)
            
            result = json.loads(response_text)
            result['raw_response'] = response_text
            result['success'] = True
            
            return result
            
        except json.JSONDecodeError as e:
            return {
                'success': False,
                'error': f'Failed to parse Gemini response as JSON: {e}',
                'raw_response': response_text if 'response_text' in dir() else None
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }


class ResumeRenderer:
    """Renders transformed content back to document format."""
    
    def render_docx(
        self,
        sections: List[Dict],
        output_path: str,
        template: Optional[ResumeTemplate] = None
    ) -> str:
        """Render sections to a DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX rendering")
        
        doc = Document()
        
        for section_data in sections:
            section_type = section_data.get('section_type', 'other')
            title = section_data.get('title', '')
            content = section_data.get('content', '')
            
            # Add section title (bold, larger)
            if title and section_type != 'header':
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(14)
            
            # Add content
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    para = doc.add_paragraph(line)
                    
                    # Style header section differently
                    if section_type == 'header':
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_path)
        return output_path
    
    def render_text(self, sections: List[Dict]) -> str:
        """Render sections to plain text."""
        output = []
        
        for section_data in sections:
            title = section_data.get('title', '')
            content = section_data.get('content', '')
            
            if title:
                output.append(f"\n{'=' * 40}")
                output.append(title.upper())
                output.append('=' * 40)
            
            output.append(content)
        
        return '\n'.join(output)


class ResumeAgent:
    """
    Main agent for resume transformation operations.
    Coordinates parsing, transformation, and rendering.
    """
    
    def __init__(self, gemini_api_key: Optional[str] = None):
        self.parser = ResumeParser()
        self.renderer = ResumeRenderer()
        self.transformer = None
        
        if gemini_api_key or os.getenv('GEMINI_API_KEY'):
            try:
                self.transformer = GeminiTransformer(gemini_api_key)
            except Exception as e:
                print(f"[!] Gemini initialization failed: {e}")
    
    def parse_resume(self, filepath: str) -> ResumeTemplate:
        """Parse a resume and return template."""
        return self.parser.parse(filepath)
    
    def transform_resume(
        self,
        input_path: str,
        transform_type: TransformType,
        output_path: Optional[str] = None,
        instructions: Optional[str] = None
    ) -> TransformResult:
        """
        Full transformation pipeline:
        1. Parse input resume
        2. Transform content via Gemini
        3. Render to output format
        """
        start_time = datetime.now()
        
        # Parse input
        try:
            template = self.parser.parse(input_path)
        except Exception as e:
            return TransformResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transform_type=transform_type.value,
                template_used=None,
                gemini_response=None,
                format_preservation_score=0.0,
                error_message=f"Parsing failed: {e}"
            )
        
        # Transform via Gemini
        if not self.transformer:
            return TransformResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transform_type=transform_type.value,
                template_used=template,
                gemini_response=None,
                format_preservation_score=0.0,
                error_message="Gemini transformer not available"
            )
        
        gemini_result = self.transformer.transform(template, transform_type, instructions)
        
        if not gemini_result.get('success'):
            return TransformResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transform_type=transform_type.value,
                template_used=template,
                gemini_response=gemini_result,
                format_preservation_score=0.0,
                error_message=gemini_result.get('error', 'Transformation failed')
            )
        
        # Render output
        if not output_path:
            input_stem = Path(input_path).stem
            output_path = f"{input_stem}_transformed.docx"
        
        try:
            self.renderer.render_docx(
                gemini_result.get('sections', []),
                output_path,
                template
            )
        except Exception as e:
            return TransformResult(
                success=False,
                input_path=input_path,
                output_path=None,
                transform_type=transform_type.value,
                template_used=template,
                gemini_response=gemini_result,
                format_preservation_score=0.0,
                error_message=f"Rendering failed: {e}"
            )
        
        processing_time = int((datetime.now() - start_time).total_seconds() * 1000)
        
        return TransformResult(
            success=True,
            input_path=input_path,
            output_path=output_path,
            transform_type=transform_type.value,
            template_used=template,
            gemini_response=gemini_result,
            format_preservation_score=gemini_result.get('format_preservation_score', 0.8),
            processing_time_ms=processing_time
        )
    
    def to_shl_message(self, action: str, task_id: str, params: Dict) -> str:
        """Format action as SHL message for agent communication."""
        params_str = json.dumps(params)
        return f"ResumeAgent|Overseer|{task_id}|{action}|{params_str}|normal"


def interactive_mode():
    """Run interactive transformation session."""
    print("\n" + "=" * 50)
    print("  JIHUB Resume Agent - Interactive Mode")
    print("=" * 50)
    
    agent = ResumeAgent()
    
    while True:
        print("\nOptions:")
        print("  1. Parse resume (extract template)")
        print("  2. Transform resume")
        print("  3. Exit")
        
        choice = input("\nSelect option (1-3): ").strip()
        
        if choice == '1':
            filepath = input("Enter resume path: ").strip()
            try:
                template = agent.parse_resume(filepath)
                print(f"\n[+] Parsed successfully!")
                print(f"    Sections found: {len(template.sections)}")
                for s in template.sections:
                    print(f"      - {s.section_type}: {s.title[:30]}...")
                
                save = input("\nSave template JSON? (y/n): ").strip().lower()
                if save == 'y':
                    out_path = filepath + ".template.json"
                    with open(out_path, 'w') as f:
                        json.dump(template.to_dict(), f, indent=2)
                    print(f"[+] Saved to: {out_path}")
                    
            except Exception as e:
                print(f"[!] Error: {e}")
        
        elif choice == '2':
            filepath = input("Enter resume path: ").strip()
            
            print("\nTransform types:")
            print("  1. content_swap - Keep format, change content")
            print("  2. style_swap - Keep content, change tone/style")
            print("  3. structure_swap - Keep content, reorder sections")
            print("  4. format_swap - Keep content, change everything else")
            
            type_choice = input("Select type (1-4): ").strip()
            type_map = {
                '1': TransformType.CONTENT_SWAP,
                '2': TransformType.STYLE_SWAP,
                '3': TransformType.STRUCTURE_SWAP,
                '4': TransformType.FORMAT_SWAP
            }
            
            transform_type = type_map.get(type_choice, TransformType.CONTENT_SWAP)
            instructions = input("Additional instructions (or blank): ").strip() or None
            
            print("\n[*] Transforming...")
            result = agent.transform_resume(filepath, transform_type, instructions=instructions)
            
            if result.success:
                print(f"[+] Success!")
                print(f"    Output: {result.output_path}")
                print(f"    Format preservation: {result.format_preservation_score:.2f}")
                print(f"    Processing time: {result.processing_time_ms}ms")
            else:
                print(f"[!] Failed: {result.error_message}")
        
        elif choice == '3':
            print("Goodbye!")
            break
        
        else:
            print("Invalid option")


def main():
    parser = argparse.ArgumentParser(description="JIHUB Resume Agent")
    subparsers = parser.add_subparsers(dest='command', help='Commands')
    
    # Parse command
    parse_parser = subparsers.add_parser('parse', help='Parse resume to template')
    parse_parser.add_argument('--input', '-i', required=True, help='Input resume path')
    parse_parser.add_argument('--output', '-o', help='Output template JSON path')
    
    # Transform command
    transform_parser = subparsers.add_parser('transform', help='Transform resume')
    transform_parser.add_argument('--input', '-i', required=True, help='Input resume path')
    transform_parser.add_argument('--output', '-o', help='Output path')
    transform_parser.add_argument('--type', '-t', choices=['content_swap', 'style_swap', 'structure_swap', 'format_swap'],
                                   default='content_swap', help='Transformation type')
    transform_parser.add_argument('--instructions', help='Additional instructions for transformation')
    
    # Interactive command
    subparsers.add_parser('interactive', help='Run interactive mode')
    
    args = parser.parse_args()
    
    if args.command == 'parse':
        agent = ResumeAgent()
        template = agent.parse_resume(args.input)
        
        output_path = args.output or (args.input + '.template.json')
        with open(output_path, 'w') as f:
            json.dump(template.to_dict(), f, indent=2)
        
        print(f"[+] Template saved to: {output_path}")
        print(f"    Sections: {len(template.sections)}")
    
    elif args.command == 'transform':
        agent = ResumeAgent()
        transform_type = TransformType(args.type)
        
        result = agent.transform_resume(
            args.input,
            transform_type,
            args.output,
            args.instructions
        )
        
        if result.success:
            print(f"[+] Transformation complete!")
            print(f"    Output: {result.output_path}")
            print(f"    Format score: {result.format_preservation_score:.2f}")
        else:
            print(f"[!] Failed: {result.error_message}")
            sys.exit(1)
    
    elif args.command == 'interactive':
        interactive_mode()
    
    else:
        parser.print_help()


if __name__ == '__main__':
    main()
